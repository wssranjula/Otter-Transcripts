"""
Document Parser for RAG Pipeline
Supports: DOCX, PDF, Excel (XLSX, XLS), TXT
Converts documents to transcript-like text format for RAG processing
"""

import io
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime
import hashlib


class DocumentParser:
    """Parse various document types into text suitable for RAG pipeline"""

    def __init__(self):
        """Initialize document parser with required libraries"""
        self.supported_extensions = ['.docx', '.pdf', '.xlsx', '.xls', '.txt']

    def is_supported(self, file_path: str) -> bool:
        """Check if file type is supported"""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions

    def parse_document(self, file_path: str, file_content: Optional[bytes] = None) -> Dict:
        """
        Parse document and return text content with metadata

        Args:
            file_path: Path to the file (used for extension detection)
            file_content: Optional bytes content (for Google Drive files)

        Returns:
            Dict with 'text', 'metadata', and 'type' keys
        """
        ext = Path(file_path).suffix.lower()

        if ext == '.docx':
            return self._parse_docx(file_path, file_content)
        elif ext == '.pdf':
            return self._parse_pdf(file_path, file_content)
        elif ext in ['.xlsx', '.xls']:
            return self._parse_excel(file_path, file_content)
        elif ext == '.txt':
            return self._parse_txt(file_path, file_content)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _parse_docx(self, file_path: str, file_content: Optional[bytes] = None) -> Dict:
        """Parse DOCX file"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")

        # Load document
        if file_content:
            doc = Document(io.BytesIO(file_content))
        else:
            doc = Document(file_path)

        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())

        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    tables_text.append(row_text)

        # Combine all text
        full_text = '\n\n'.join(paragraphs)
        if tables_text:
            full_text += '\n\n[TABLES]\n' + '\n'.join(tables_text)

        # Extract metadata
        core_props = doc.core_properties
        metadata = {
            'title': core_props.title or Path(file_path).stem,
            'author': core_props.author or 'Unknown',
            'created': core_props.created.isoformat() if core_props.created else None,
            'modified': core_props.modified.isoformat() if core_props.modified else None,
            'source_file': Path(file_path).name,
            'source_type': 'docx',
            'paragraph_count': len(paragraphs),
            'table_count': len(doc.tables)
        }

        return {
            'text': full_text,
            'metadata': metadata,
            'type': 'document'
        }

    def _parse_pdf(self, file_path: str, file_content: Optional[bytes] = None) -> Dict:
        """Parse PDF file"""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("PyPDF2 is required. Install with: pip install PyPDF2")

        # Load PDF
        if file_content:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        else:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)

        # Extract text from all pages
        pages_text = []
        for page_num, page in enumerate(pdf_reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                pages_text.append(f"[Page {page_num}]\n{text.strip()}")

        full_text = '\n\n'.join(pages_text)

        # Extract metadata
        metadata_dict = pdf_reader.metadata if pdf_reader.metadata else {}
        metadata = {
            'title': metadata_dict.get('/Title', Path(file_path).stem),
            'author': metadata_dict.get('/Author', 'Unknown'),
            'created': metadata_dict.get('/CreationDate', None),
            'modified': metadata_dict.get('/ModDate', None),
            'source_file': Path(file_path).name,
            'source_type': 'pdf',
            'page_count': len(pdf_reader.pages)
        }

        return {
            'text': full_text,
            'metadata': metadata,
            'type': 'document'
        }

    def _parse_excel(self, file_path: str, file_content: Optional[bytes] = None) -> Dict:
        """Parse Excel file"""
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("pandas and openpyxl are required. Install with: pip install pandas openpyxl")

        # Load Excel file
        if file_content:
            excel_file = pd.ExcelFile(io.BytesIO(file_content))
        else:
            excel_file = pd.ExcelFile(file_path)

        # Process each sheet
        sheets_text = []
        total_rows = 0

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)

            if df.empty:
                continue

            # Convert dataframe to readable text
            sheet_text = f"[Sheet: {sheet_name}]\n"
            sheet_text += f"Columns: {', '.join(str(col) for col in df.columns)}\n\n"

            # Convert rows to text
            for idx, row in df.iterrows():
                row_text = ' | '.join(f"{col}: {val}" for col, val in row.items() if pd.notna(val))
                if row_text:
                    sheet_text += f"Row {idx + 1}: {row_text}\n"

            sheets_text.append(sheet_text)
            total_rows += len(df)

        full_text = '\n\n'.join(sheets_text)

        # Metadata
        metadata = {
            'title': Path(file_path).stem,
            'author': 'Unknown',
            'created': None,
            'modified': None,
            'source_file': Path(file_path).name,
            'source_type': 'excel',
            'sheet_count': len(excel_file.sheet_names),
            'total_rows': total_rows,
            'sheets': excel_file.sheet_names
        }

        return {
            'text': full_text,
            'metadata': metadata,
            'type': 'spreadsheet'
        }

    def _parse_txt(self, file_path: str, file_content: Optional[bytes] = None) -> Dict:
        """Parse plain text file"""
        # Load text content
        if file_content:
            text = file_content.decode('utf-8', errors='ignore')
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()

        # Basic metadata
        metadata = {
            'title': Path(file_path).stem,
            'author': 'Unknown',
            'created': None,
            'modified': None,
            'source_file': Path(file_path).name,
            'source_type': 'txt',
            'character_count': len(text),
            'line_count': len(text.splitlines())
        }

        return {
            'text': text.strip(),
            'metadata': metadata,
            'type': 'document'
        }

    def convert_to_transcript_format(self, parsed_doc: Dict, timestamp: Optional[str] = None) -> str:
        """
        Convert parsed document to transcript-like format for RAG pipeline

        This format mimics transcript structure so it can be processed
        by the existing RAG pipeline
        """
        metadata = parsed_doc['metadata']
        text = parsed_doc['text']

        # Create timestamp
        if not timestamp:
            if metadata.get('created'):
                timestamp = metadata['created']
            else:
                timestamp = datetime.now().isoformat()

        # Format as pseudo-transcript
        transcript = f"""Document: {metadata['title']}
Source: {metadata['source_file']}
Type: {metadata['source_type'].upper()}
Author: {metadata['author']}
Date: {timestamp}

=== CONTENT ===

{text}
"""

        return transcript

    def save_as_transcript(self, file_path: str, output_dir: str,
                          file_content: Optional[bytes] = None) -> str:
        """
        Parse document and save as transcript file

        Returns:
            Path to saved transcript file
        """
        # Parse document
        parsed = self.parse_document(file_path, file_content)

        # Convert to transcript format
        transcript_text = self.convert_to_transcript_format(parsed)

        # Generate output filename
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        base_name = Path(file_path).stem
        # Add hash to prevent conflicts
        hash_suffix = hashlib.md5(file_path.encode()).hexdigest()[:6]
        output_file = output_dir / f"{base_name}_{hash_suffix}.txt"

        # Save
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(transcript_text)

        return str(output_file)


def main():
    """Test document parser"""
    parser = DocumentParser()

    print("="*70)
    print("DOCUMENT PARSER TEST")
    print("="*70)
    print(f"\nSupported formats: {', '.join(parser.supported_extensions)}")
    print("\nUsage:")
    print("  parser = DocumentParser()")
    print("  parsed = parser.parse_document('document.docx')")
    print("  transcript = parser.convert_to_transcript_format(parsed)")
    print("\nOr save directly:")
    print("  output_file = parser.save_as_transcript('doc.docx', 'transcripts/')")
    print("="*70)


if __name__ == "__main__":
    main()
